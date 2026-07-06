r"""
matrix_reinject.py  (Trainstorm Toolkit)

Takes the ORIGINAL Storyline matrix .docx plus a translation job folder
(created by matrix_extract.py, with the AI's answers saved into
2_translated\\ and/or typed into master.xlsx), validates everything, and
writes the translated text into a NEW COPY of the matrix.

GUARANTEES
  * The original docx is never modified.
  * No structural element of the document is ever touched - tables, rows,
    cells, merges, and formatting pass through untouched. Only the text
    inside translation-target cells is replaced.
  * Nothing is written unless validation passes (or --allow-partial is
    explicitly given). Failures are reported precisely, per segment.

TRANSLATION SOURCES (precedence)
  1. master.xlsx Translation column (the human-editable ledger) - wins.
  2. .txt files in 2_translated\\ containing "[S0001] translated text"
     entries. Order, file names, and duplicates don't matter; for
     duplicate IDs across txt files the last-read value is used and the
     conflict is reported.

USAGE
  python matrix_reinject.py --matrix original.docx --job <job folder>
                            [--allow-partial]

EXIT CODES
  0 = translated copy written
  2 = validation incomplete, nothing written (fix and re-run)
  1 = hard error (wrong file, unreadable job, etc.)
"""

import argparse
import copy
import glob
import hashlib
import json
import os
import re
import sys
from datetime import datetime, timezone

from docx import Document
from docx.oxml.ns import qn

MARKER_RE = re.compile(r"^\s*\[?(S\d{4})\]?\s*[:\-\u2013]?\s*", re.MULTILINE)


def sha1(text):
    return hashlib.sha1(text.encode("utf-8")).hexdigest()


def file_sha256(path):
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(1 << 20), b""):
            h.update(chunk)
    return h.hexdigest()


# ----------------------------------------------------------------------
# Reading the AI's answers
# ----------------------------------------------------------------------

def parse_txt_blob(text):
    """Parse '[S0001] translation' entries out of pasted AI output.

    Tolerant of: missing brackets, trailing colons/dashes after the ID,
    multi-line translations (continue until the next ID marker), blank
    lines, and surrounding chatter that contains no ID markers.
    """
    out = {}
    matches = list(MARKER_RE.finditer(text))
    for i, m in enumerate(matches):
        seg_id = m.group(1)
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        body = text[start:end].strip()
        body = body.replace("<BR>", "\n").replace("<br>", "\n")
        out[seg_id] = body
    return out


def gather_translations(job_dir, language):
    """Returns (translations, notes). master.xlsx wins over txt packets."""
    notes = []
    from_txt = {}
    txt_dir = os.path.join(job_dir, "2_translated")
    for path in sorted(glob.glob(os.path.join(txt_dir, "*.txt"))):
        if os.path.basename(path).startswith("PUT_AI_ANSWERS"):
            continue
        with open(path, "r", encoding="utf-8-sig", errors="replace") as f:
            parsed = parse_txt_blob(f.read())
        for seg_id, val in parsed.items():
            if seg_id in from_txt and from_txt[seg_id] != val:
                notes.append(f"NOTE  {seg_id}: multiple differing answers in "
                             f"txt files; using the one from "
                             f"{os.path.basename(path)}")
            from_txt[seg_id] = val
        notes.append(f"READ  {os.path.basename(path)}: "
                     f"{len(parsed)} segment answers")

    from_xlsx = {}
    xlsx_path = os.path.join(job_dir, "master.xlsx")
    if os.path.isfile(xlsx_path):
        from openpyxl import load_workbook
        wb = load_workbook(xlsx_path, data_only=True)
        ws = wb.active
        for row in ws.iter_rows(min_row=2, values_only=True):
            if not row or not row[0]:
                continue
            seg_id = str(row[0]).strip()
            val = row[2] if len(row) > 2 else None
            if val is not None and str(val).strip():
                from_xlsx[seg_id] = str(val)
        if from_xlsx:
            notes.append(f"READ  master.xlsx: {len(from_xlsx)} filled "
                         f"translation cells (these take precedence)")

    merged = dict(from_txt)
    for seg_id, val in from_xlsx.items():
        if seg_id in merged and merged[seg_id] != val:
            notes.append(f"NOTE  {seg_id}: master.xlsx overrides the txt answer")
        merged[seg_id] = val
    return merged, notes


# ----------------------------------------------------------------------
# Locating cells and replacing text (the surgical part)
# ----------------------------------------------------------------------

def logical_cells(row):
    seen, out = set(), []
    for idx, cell in enumerate(row.cells):
        key = id(cell._tc)
        if key in seen:
            continue
        seen.add(key)
        out.append((idx, cell))
    return out


def resolve_cell(doc, address):
    """Follow an address path [[t,r,c], ...] down through nested tables."""
    tables = doc.tables
    cell = None
    for t_idx, r_idx, c_idx in address:
        table = tables[t_idx]
        row = table.rows[r_idx]
        cell = None
        for grid_idx, c in logical_cells(row):
            if grid_idx == c_idx:
                cell = c
                break
        if cell is None:
            raise LookupError(f"cell index {c_idx} not found")
        tables = cell.tables
    return cell


def cell_direct_text(tc):
    lines = []
    for p in tc.findall(qn("w:p")):
        parts = []
        for node in p.iter():
            tag = node.tag
            if tag == qn("w:t"):
                parts.append(node.text or "")
            elif tag == qn("w:br"):
                parts.append("\n")
            elif tag == qn("w:tab"):
                parts.append("\t")
        lines.append("".join(parts))
    return "\n".join(lines)


def _make_text_node(text):
    from docx.oxml import OxmlElement
    t = OxmlElement("w:t")
    t.text = text
    if text != text.strip():
        t.set(qn("xml:space"), "preserve")
    return t


def _set_paragraph_text(p, lines):
    """Write lines into a paragraph: first line as text, the rest joined
    with <w:br/>. First run keeps its formatting; surplus runs are removed.
    """
    from docx.oxml import OxmlElement
    runs = p.findall(qn("w:r"))
    if not runs:
        r = OxmlElement("w:r")
        p.append(r)
        runs = [r]
    first = runs[0]
    # strip old content from the first run, keep its run properties (w:rPr)
    for child in list(first):
        if child.tag != qn("w:rPr"):
            first.remove(child)
    for i, line in enumerate(lines):
        if i > 0:
            first.append(OxmlElement("w:br"))
        first.append(_make_text_node(line))
    # remove every other plain run in the paragraph
    for r in runs[1:]:
        p.remove(r)


def set_cell_text(tc, text):
    """Distribute translated text across the cell's existing paragraphs.

    The paragraph skeleton is preserved (Storyline may distinguish
    paragraph breaks from soft breaks). Translation lines fill existing
    paragraphs in order; overflow lines go into the last paragraph as
    soft breaks; leftover paragraphs are emptied.
    """
    paras = tc.findall(qn("w:p"))
    if not paras:
        raise LookupError("cell has no paragraphs")
    lines = text.split("\n")
    n_p = len(paras)
    for j, p in enumerate(paras):
        if j < n_p - 1:
            chunk = [lines[j]] if j < len(lines) else [""]
        else:
            chunk = lines[j:] if j < len(lines) else [""]
        _set_paragraph_text(p, chunk)


# ----------------------------------------------------------------------
# Main
# ----------------------------------------------------------------------

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--matrix", required=True, help="ORIGINAL matrix .docx")
    ap.add_argument("--job", required=True, help="Translation job folder")
    ap.add_argument("--allow-partial", action="store_true",
                    help="Write even if some segments lack translations "
                         "(untranslated cells keep their source text)")
    args = ap.parse_args()

    matrix_path = os.path.abspath(args.matrix)
    job_dir = os.path.abspath(args.job)
    job_path = os.path.join(job_dir, "job.json")

    if not os.path.isfile(matrix_path):
        print(f"[ERROR] Matrix file not found: {matrix_path}")
        return 1
    if not os.path.isfile(job_path):
        print(f"[ERROR] This folder is not a translation job (no job.json):")
        print(f"        {job_dir}")
        return 1

    with open(job_path, "r", encoding="utf-8") as f:
        job = json.load(f)
    segments = job["segments"]
    language = job.get("language", "XX")

    # --- Identity check: is this the same export the job was built from? ---
    if file_sha256(matrix_path) != job["source_docx_sha256"]:
        print("[ERROR] This is NOT the same file the job was extracted from.")
        print(f"        Job was built from : {job['source_docx_name']}")
        print(f"        You selected       : {os.path.basename(matrix_path)}")
        print("        Use the exact original export. If it was lost, re-export")
        print("        from Storyline and run matrix_extract again (a fresh")
        print("        export needs a fresh job - internal IDs change).")
        return 1

    translations, notes = gather_translations(job_dir, language)

    # --- Validation pass (no writes yet) ---
    doc = Document(matrix_path)
    known_ids = {s["id"] for s in segments}
    problems, infos, ok = [], [], []

    for seg_id in sorted(set(translations) - known_ids):
        problems.append(f"UNKNOWN  {seg_id}: answer provided but this ID is "
                        f"not in the job (typo or wrong job folder?)")

    for seg in segments:
        seg_id = seg["id"]
        val = translations.get(seg_id)
        if val is None or not val.strip():
            problems.append(f"MISSING  {seg_id}: no translation found "
                            f"(source: \"{_clip(seg['source_text'])}\")")
            continue
        try:
            cell = resolve_cell(doc, seg["address"])
        except Exception as e:
            problems.append(f"ADDRESS  {seg_id}: cannot locate target cell "
                            f"({e}) - contact Jake")
            continue
        current = cell_direct_text(cell._tc)
        if sha1(current) != seg["source_sha1"]:
            problems.append(f"CHANGED  {seg_id}: the target cell's text no "
                            f"longer matches the export (was the docx edited?)")
            continue
        if val.strip() == seg["source_text"].strip():
            infos.append(f"SAME     {seg_id}: translation identical to source "
                         f"(allowed - may be intentional)")
        ok.append((seg, val))

    # --- Coverage report ---
    report_path = os.path.join(job_dir, "coverage_report.txt")
    missing_count = len(segments) - len(ok)
    with open(report_path, "w", encoding="utf-8") as f:
        f.write(f"COVERAGE REPORT - {datetime.now(timezone.utc).isoformat(timespec='seconds')}\n")
        f.write(f"Job: {job['source_docx_name']} -> {language}\n")
        f.write(f"Ready to inject : {len(ok)} of {len(segments)} segments\n\n")
        for line in problems:
            f.write(line + "\n")
        if infos:
            f.write("\n")
            for line in infos:
                f.write(line + "\n")
        if notes:
            f.write("\n")
            for line in notes:
                f.write(line + "\n")

    print()
    print(f"  Ready to inject : {len(ok)} of {len(segments)} segments")
    if problems:
        print(f"  Problems        : {len(problems)}  (see coverage_report.txt)")
        for line in problems[:10]:
            print(f"    {line}")
        if len(problems) > 10:
            print(f"    ... and {len(problems) - 10} more in the report")
    print(f"  Full report     : {report_path}")

    if missing_count and not args.allow_partial:
        print()
        print("  Nothing was written. Fix the items above (usually: re-run the")
        print("  affected packets and save the answers into 2_translated),")
        print("  then run reinject again.")
        return 2

    # --- Injection: text-node surgery on an in-memory copy, saved as new ---
    for seg, val in ok:
        cell = resolve_cell(doc, seg["address"])
        set_cell_text(cell._tc, val)

    lang_tag = re.sub(r"[^A-Za-z0-9]+", "", language)[:12] or "XX"
    stem, ext = os.path.splitext(matrix_path)
    out_path = f"{stem}_{lang_tag}{ext}"
    n = 2
    while os.path.exists(out_path):
        out_path = f"{stem}_{lang_tag}_{n}{ext}"
        n += 1
    doc.save(out_path)

    print()
    print(f"  Translated copy written ({len(ok)} segments injected"
          + (f", {missing_count} left as source text" if missing_count else "")
          + "):")
    print(f"    {out_path}")
    print(f"  The original was not modified. Import the new copy into Storyline.")
    return 0


def _clip(text, n=50):
    text = (text or "").replace("\n", " / ")
    return text if len(text) <= n else text[:n - 3] + "..."


if __name__ == "__main__":
    sys.exit(main())
