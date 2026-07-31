#!/usr/bin/env python3
"""
matrix_tools.py — safely read and edit an Articulate Storyline translation matrix
exported to Word (.docx), so reviewer edits re-import cleanly.

The Word "Export to Translation" file is a .docx whose SECOND table has columns:
    [0] ID (locked)   [1] Type   [2] Source Text   [3] Translation
Storyline re-imports by reading the *Translation* column. Each on-screen paragraph
is its own row; the Source Text column repeats the whole text box (with 1/2/3
markers) for every row of that box, while the Translation cell holds one paragraph.

Golden rule: change ONLY the character data inside existing runs of the Translation
cell. Never add/remove/reorder rows, never touch the ID column, and preserve run
formatting (bold etc.) so it survives import.

Commands
--------
dump    Print every content row: index, Type, Source, Translation, run formatting.
        Use this to match reviewer comments to rows by Source Text.
          python matrix_tools.py dump FILE.docx [--grep TEXT] [--runs]

apply   Apply edits from a JSON file and verify integrity. Each edit targets a row
        (by index, or by an unambiguous source/translation anchor) and does one of:
          replace  : swap a substring inside the matched run (preserves its format)
          append   : add text to the end of the cell's last run (new paragraph via \\n)
          set      : replace the whole cell's (first run) text
          clear    : empty the cell's runs (removes a paragraph; keeps the row)
          python matrix_tools.py apply FILE.docx edits.json OUT.docx

verify  Re-check an edited file against the original: row count, ID column, and
        Source column must be unchanged (only Translation may differ).
          python matrix_tools.py verify ORIGINAL.docx EDITED.docx

edits.json schema (a list):
[
  {"row": 17, "mode": "set", "text": "New paragraph text"},
  {"row": 478, "mode": "replace", "find": "CAPA for Medical Devices ", "text": "CAPA Management "},
  {"row": 629, "mode": "append", "text": "Note: ..."},
  {"source_anchor": "On the next screen", "mode": "clear"}
]
Prefer "row" (from dump) when you can; "source_anchor" matches the first content row
whose Source Text contains that string (must be unique enough).
"""
import sys, json
from docx import Document

TRANS_COL = 3
SRC_COL = 2
ID_COL = 0


def _table(doc):
    # the translation grid is the table with an "ID"/"Translation" header
    for t in doc.tables:
        hdr = [c.text.strip().lower() for c in t.rows[0].cells]
        if any("translation" in h for h in hdr):
            return t
    return doc.tables[-1]


def _runs(cell):
    return [r for p in cell.paragraphs for r in p.runs]


def cmd_dump(path, grep=None, show_runs=False):
    t = _table(Document(path))
    for i in range(1, len(t.rows)):
        src = t.rows[i].cells[SRC_COL].text
        trans = t.rows[i].cells[TRANS_COL].text
        if grep and grep.lower() not in (src + trans).lower():
            continue
        typ = t.rows[i].cells[1].text
        print(f"[{i}] Type={typ!r}")
        print(f"    SRC  : {src[:120]!r}")
        print(f"    TRANS: {trans!r}")
        if show_runs:
            rs = [(("B" if r.bold else "") + ("I" if r.italic else "") or "-", r.text)
                  for r in _runs(t.rows[i].cells[TRANS_COL])]
            print(f"    RUNS : {rs}")


def _find_row(t, edit):
    if "row" in edit:
        return edit["row"]
    # Match on the Translation cell (the individual paragraph), NOT Source Text:
    # the Source column repeats the whole text box on every row of that box, so
    # anchoring there matches every paragraph. Translation holds one paragraph.
    anchor = edit.get("anchor", edit.get("source_anchor"))
    hits = [i for i in range(1, len(t.rows)) if anchor in t.rows[i].cells[TRANS_COL].text]
    if len(hits) != 1:
        raise SystemExit(f"anchor {anchor!r} matched {len(hits)} rows {hits}; "
                         f"narrow it or use an explicit row index from `dump`.")
    return hits[0]


def _apply_one(cell, edit):
    runs = _runs(cell)
    mode = edit["mode"]
    if mode == "replace":
        find = edit["find"]
        for r in runs:
            if find in r.text:
                r.text = r.text.replace(find, edit["text"])
                return
        raise SystemExit(f"replace: {find!r} not found in Translation cell")
    if mode == "append":
        if not runs:
            raise SystemExit("append: cell has no runs to append to")
        runs[-1].text = runs[-1].text + edit["text"]
        return
    if mode == "set":
        if not runs:
            raise SystemExit("set: cell has no runs")
        runs[0].text = edit["text"]
        for r in runs[1:]:
            if r.text.strip():
                raise SystemExit("set: multiple non-empty runs; use replace instead")
        return
    if mode == "clear":
        for r in runs:
            r.text = ""
        return
    raise SystemExit(f"unknown mode {mode!r}")


def cmd_apply(path, edits_path, out_path):
    doc = Document(path)
    t = _table(doc)
    edits = json.load(open(edits_path))
    for e in edits:
        row = _find_row(t, e)
        _apply_one(t.rows[row].cells[TRANS_COL], e)
        print(f"applied {e['mode']} -> row {row}")
    doc.save(out_path)
    ok = _verify(path, out_path)
    print("SAVED", out_path, "| integrity:", "OK" if ok else "*** FAILED ***")
    if not ok:
        sys.exit(1)


def _verify(orig_path, edited_path):
    a = _table(Document(orig_path))
    b = _table(Document(edited_path))
    if len(a.rows) != len(b.rows):
        print(f"  row count changed: {len(a.rows)} -> {len(b.rows)}")
        return False
    id_mis = sum(1 for i in range(len(a.rows)) if a.rows[i].cells[ID_COL].text != b.rows[i].cells[ID_COL].text)
    src_mis = sum(1 for i in range(len(a.rows)) if a.rows[i].cells[SRC_COL].text != b.rows[i].cells[SRC_COL].text)
    changed = sum(1 for i in range(len(a.rows)) if a.rows[i].cells[TRANS_COL].text != b.rows[i].cells[TRANS_COL].text)
    print(f"  rows={len(b.rows)} id_mismatch={id_mis} source_mismatch={src_mis} translation_changed={changed}")
    return id_mis == 0 and src_mis == 0


def cmd_verify(orig, edited):
    sys.exit(0 if _verify(orig, edited) else 1)


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print(__doc__); sys.exit(1)
    cmd = sys.argv[1]
    if cmd == "dump":
        grep = None; runs = "--runs" in sys.argv
        if "--grep" in sys.argv:
            grep = sys.argv[sys.argv.index("--grep") + 1]
        cmd_dump(sys.argv[2], grep, runs)
    elif cmd == "apply":
        cmd_apply(sys.argv[2], sys.argv[3], sys.argv[4])
    elif cmd == "verify":
        cmd_verify(sys.argv[2], sys.argv[3])
    else:
        print(__doc__); sys.exit(1)