r"""
matrix_extract.py  (Trainstorm Toolkit)

Reads a Storyline translation matrix (.docx export), finds every translatable
text segment, and produces a "translation job" folder containing:

  job.json            - the machine record: every segment, its address in the
                        docx, and a fingerprint of the source text
  master.xlsx         - the human-readable ledger (ID / Source / Translation)
  1_packets\          - small numbered text files, each with the translation
                        prompt baked in, sized for weaker LLMs
  2_translated\       - EMPTY. The user saves the LLM's responses here.
  INSTRUCTIONS.txt    - step-by-step guide for the person running the job
  structure_report.txt- what the extractor found, skipped, and why

DESIGN RULES
  * The docx is READ ONLY here. Nothing ever modifies it.
  * Tolerant reader, loud failer: anything structurally ambiguous is
    reported and skipped, never guessed at.
  * Storyline's export duplicates the source text into the translation
    column. We detect translation targets by that duplication.

USAGE
  python matrix_extract.py --input matrix.docx --language Japanese
                           [--packet-size 12] [--output <folder>]
"""

import argparse
import hashlib
import json
import os
import re
import sys
from datetime import datetime, timezone

from docx import Document
from docx.oxml.ns import qn

TOOL_VERSION = "1.0"
DEFAULT_PACKET_SIZE = 12


# ----------------------------------------------------------------------
# Low-level text reading (direct XML walk so nested tables don't bleed in)
# ----------------------------------------------------------------------

def cell_direct_text(tc):
    """Text of a table cell from its DIRECT paragraphs only.

    Nested tables inside the cell are excluded (they are extracted
    separately with their own addresses). Line breaks (<w:br/>) and
    paragraph boundaries both become '\n'. Tabs become '\t'.
    """
    lines = []
    for p in tc.findall(qn("w:p")):
        parts = []
        for node in p.iter():
            tag = node.tag
            if tag == qn("w:t"):
                parts.append(node.text or "")
            elif tag == qn("w:br"):
                parts.append("\n")
            elif tag == qn("w:tab"):
                parts.append("\t")
        lines.append("".join(parts))
    return "\n".join(lines)


def norm(text):
    """Normalization used ONLY for source==translation duplicate detection."""
    return re.sub(r"\s+", " ", text or "").strip()


def sha1(text):
    return hashlib.sha1(text.encode("utf-8")).hexdigest()


def file_sha256(path):
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(1 << 20), b""):
            h.update(chunk)
    return h.hexdigest()


# ----------------------------------------------------------------------
# Table walking
# ----------------------------------------------------------------------

def logical_cells(row):
    """Cells of a row with merged duplicates removed.

    python-docx repeats the same underlying <w:tc> for horizontally merged
    cells; we keep the first occurrence and remember its grid index so the
    reinjector (which dedupes the same way) lands on the same cell.
    """
    seen = set()
    out = []
    for idx, cell in enumerate(row.cells):
        key = id(cell._tc)
        if key in seen:
            continue
        seen.add(key)
        out.append((idx, cell))
    return out


def walk_tables(tables, base_path, report, segments, counters):
    """Recursively extract segments from a list of tables.

    base_path is a list of [table, row, cell] triples locating a nested
    table's host cell; top-level tables have base_path == [].
    """
    for t_idx, table in enumerate(tables):
        rows = list(table.rows)
        # First pass: gather row cell texts and find duplicate pattern
        row_data = []
        for r_idx, row in enumerate(rows):
            cells = logical_cells(row)
            texts = [cell_direct_text(c._tc) for _, c in cells]
            row_data.append((r_idx, cells, texts))

        candidates = []   # rows matching the source==translation pattern
        nonempty_last = 0
        for r_idx, cells, texts in row_data:
            if len(cells) >= 2 and norm(texts[-1]):
                nonempty_last += 1
                if norm(texts[-1]) == norm(texts[-2]):
                    candidates.append((r_idx, cells, texts))

        table_label = _path_label(base_path, t_idx)

        if not candidates:
            report.append(
                f"SKIPPED  table {table_label}: no rows where the last column "
                f"duplicates the previous column ({len(rows)} rows). "
                f"Not a translation table, or an unexpected layout."
            )
        else:
            ratio = len(candidates) / max(nonempty_last, 1)
            if ratio < 0.6:
                report.append(
                    f"SKIPPED  table {table_label}: only {len(candidates)} of "
                    f"{nonempty_last} rows match the translation pattern "
                    f"({ratio:.0%}). Ambiguous layout - review structure_report "
                    f"and contact Jake before trusting this table."
                )
            else:
                matched_rows = {r for r, _, _ in candidates}
                for r_idx, cells, texts in candidates:
                    grid_idx = cells[-1][0]          # target = last column
                    counters["seg"] += 1
                    seg_id = f"S{counters['seg']:04d}"
                    segments.append({
                        "id": seg_id,
                        "address": base_path + [[t_idx, r_idx, grid_idx]],
                        "source_text": texts[-2],
                        "source_sha1": sha1(texts[-1]),  # fingerprint of the
                        # CURRENT target-cell text (== source in a fresh export)
                    })
                skipped = [r for r, _, _ in row_data
                           if r not in matched_rows and any(norm(t) for t in row_data[r][2])]
                report.append(
                    f"OK       table {table_label}: {len(candidates)} segments "
                    f"extracted" + (f"; rows not matching pattern (headers/labels): "
                                    f"{skipped}" if skipped else "")
                )

        # Recurse into nested tables regardless of parent qualification
        for r_idx, cells, _texts in row_data:
            for grid_idx, cell in cells:
                nested = cell.tables
                if nested:
                    walk_tables(nested,
                                base_path + [[t_idx, r_idx, grid_idx]],
                                report, segments, counters)


def _path_label(base_path, t_idx):
    if not base_path:
        return f"#{t_idx}"
    outer = ".".join(f"t{t}r{r}c{c}" for t, r, c in base_path)
    return f"{outer}.#{t_idx} (nested)"


# ----------------------------------------------------------------------
# Outputs
# ----------------------------------------------------------------------

PACKET_PROMPT = """\
You are translating eLearning course text exported from Articulate Storyline.
Translate each segment below into {language}.

RULES - follow these exactly:
1. Output ONE entry per segment, in exactly this format:
   [S0001] translated text here
2. Keep each bracketed ID exactly as given. Never renumber, reorder,
   merge, split, or skip a segment.
3. Output ONLY the ID lines. No commentary, headers, notes, or markdown.
4. If a segment contains the marker <BR>, keep a <BR> marker at the
   equivalent position in your translation. Do not add <BR> markers that
   were not in the source.
5. Do not translate placeholder codes or variable references such as
   %UserName% - copy them through unchanged.
6. If a segment should stay untranslated (a product name, a number, a
   button code), repeat it unchanged after its ID.

SEGMENTS - packet {packet_no} of {packet_total} ({count} segments):

"""


def write_packets(job_dir, segments, language, packet_size):
    pdir = os.path.join(job_dir, "1_packets")
    os.makedirs(pdir, exist_ok=True)
    chunks = [segments[i:i + packet_size]
              for i in range(0, len(segments), packet_size)]
    for i, chunk in enumerate(chunks, 1):
        body = PACKET_PROMPT.format(language=language, packet_no=i,
                                    packet_total=len(chunks), count=len(chunk))
        lines = []
        for seg in chunk:
            flat = seg["source_text"].replace("\n", "<BR>")
            lines.append(f"[{seg['id']}] {flat}")
        path = os.path.join(pdir, f"packet_{i:02d}.txt")
        with open(path, "w", encoding="utf-8") as f:
            f.write(body + "\n".join(lines) + "\n")
    return len(chunks)


def write_master_xlsx(job_dir, segments, language):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment

    wb = Workbook()
    ws = wb.active
    ws.title = "Translation"
    headers = ["ID", "Source Text", f"Translation ({language})"]
    ws.append(headers)
    hdr_fill = PatternFill("solid", start_color="4D4D4F")
    for c in ws[1]:
        c.font = Font(bold=True, color="FFFFFF", name="Arial")
        c.fill = hdr_fill
    for seg in segments:
        ws.append([seg["id"], seg["source_text"], ""])
    wrap = Alignment(wrap_text=True, vertical="top")
    for row in ws.iter_rows(min_row=2):
        for c in row:
            c.alignment = wrap
            c.font = Font(name="Arial")
    ws.column_dimensions["A"].width = 10
    ws.column_dimensions["B"].width = 60
    ws.column_dimensions["C"].width = 60
    ws.freeze_panes = "A2"
    path = os.path.join(job_dir, "master.xlsx")
    wb.save(path)


INSTRUCTIONS = """\
TRANSLATION JOB - WHAT TO DO
============================================================
Course matrix : {matrix_name}
Target language : {language}
Segments : {n_segments}   Packets : {n_packets}
Created : {created}

STEP 1 - TRANSLATE THE PACKETS
  Open the "1_packets" folder. For each packet file:
    a. Open it, Select All, Copy.
    b. Paste the WHOLE thing into Copilot / Studio and send.
       (The instructions for the AI are already inside the packet.)
    c. Copy the AI's complete answer.
    d. Paste it into a new text file saved in the "2_translated"
       folder. Any file name is fine. One file per packet, or all
       answers in one big file - both work.
  Do the packets in any order. If an answer looks wrong, just run
  that packet again and save the new answer - the tool keeps the
  best-validated version.

  ALTERNATIVE: you can instead type or paste translations directly
  into the Translation column of master.xlsx. Anything filled in
  there wins over the packet answers.

STEP 2 - REINJECT
  Double-click matrix_reinject.bat in the Trainstorm Toolkit folder.
    a. When asked, pick the ORIGINAL matrix .docx (the exact file
       you extracted from - do not re-export from Storyline).
    b. When asked, pick THIS job folder.
  The tool checks every translation and tells you exactly what is
  missing or malformed. If anything is incomplete, fix or re-run
  just those packets, drop the answers in "2_translated", and run
  reinject again.

STEP 3 - IMPORT
  When validation passes, the tool writes a NEW translated copy of
  the matrix next to the original (the original is never touched).
  Import that new file into Storyline.

RULES THAT KEEP THIS SAFE
  * Never edit the original .docx by hand or with an AI.
  * Never re-export the matrix from Storyline mid-job (it changes
    internal IDs). Finish the round trip on one export.
  * Don't edit the ID column anywhere.
Questions? Contact Jake.
"""


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--input", required=True, help="Storyline matrix .docx")
    ap.add_argument("--language", required=True, help="Target language, e.g. Japanese")
    ap.add_argument("--packet-size", type=int, default=DEFAULT_PACKET_SIZE)
    ap.add_argument("--output", help="Job folder (default: next to the docx)")
    args = ap.parse_args()

    matrix_path = os.path.abspath(args.input)
    if not os.path.isfile(matrix_path):
        print(f"[ERROR] File not found: {matrix_path}")
        return 1

    stem = os.path.splitext(os.path.basename(matrix_path))[0]
    lang_tag = re.sub(r"[^A-Za-z0-9]+", "", args.language)[:12] or "XX"
    job_dir = os.path.abspath(args.output) if args.output else os.path.join(
        os.path.dirname(matrix_path), f"{stem}_{lang_tag}_translation_job")
    os.makedirs(job_dir, exist_ok=True)
    os.makedirs(os.path.join(job_dir, "2_translated"), exist_ok=True)

    print(f"Reading matrix: {matrix_path}")
    doc = Document(matrix_path)

    report, segments = [], []
    walk_tables(doc.tables, [], report, segments, {"seg": 0})

    # Loud failure if nothing usable was found
    if not segments:
        rpt = os.path.join(job_dir, "structure_report.txt")
        with open(rpt, "w", encoding="utf-8") as f:
            f.write("NO TRANSLATABLE SEGMENTS FOUND\n\n" + "\n".join(report))
        print("[ERROR] No translatable segments found in this document.")
        print(f"        Details written to: {rpt}")
        print("        This usually means the layout is not the expected")
        print("        Storyline translation export. Send the docx and the")
        print("        report to Jake.")
        return 1

    job = {
        "tool": "matrix_extract", "tool_version": TOOL_VERSION,
        "created": datetime.now(timezone.utc).isoformat(timespec="seconds"),
        "source_docx_name": os.path.basename(matrix_path),
        "source_docx_sha256": file_sha256(matrix_path),
        "language": args.language,
        "packet_size": args.packet_size,
        "segments": segments,
    }
    with open(os.path.join(job_dir, "job.json"), "w", encoding="utf-8") as f:
        json.dump(job, f, ensure_ascii=False, indent=1)

    n_packets = write_packets(job_dir, segments, args.language, args.packet_size)
    write_master_xlsx(job_dir, segments, args.language)

    with open(os.path.join(job_dir, "structure_report.txt"), "w", encoding="utf-8") as f:
        f.write(f"STRUCTURE REPORT - {os.path.basename(matrix_path)}\n")
        f.write(f"{len(segments)} segments extracted\n\n" + "\n".join(report) + "\n")

    with open(os.path.join(job_dir, "INSTRUCTIONS.txt"), "w", encoding="utf-8") as f:
        f.write(INSTRUCTIONS.format(
            matrix_name=os.path.basename(matrix_path), language=args.language,
            n_segments=len(segments), n_packets=n_packets,
            created=job["created"]))

    with open(os.path.join(job_dir, "2_translated",
                           "PUT_AI_ANSWERS_HERE.txt"), "w", encoding="utf-8") as f:
        f.write("Save the AI's answers in this folder as .txt files.\n"
                "Any file names are fine. See INSTRUCTIONS.txt one level up.\n")

    print()
    print(f"  Segments extracted : {len(segments)}")
    print(f"  Packets created    : {n_packets} (x{args.packet_size} segments)")
    print(f"  Job folder         : {job_dir}")
    print()
    print("  Next: open INSTRUCTIONS.txt in the job folder and follow Step 1.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
